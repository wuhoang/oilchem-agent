"""
实验报告生成服务（M7 增强）。

根据实验数据生成：
- Word 报告（实验信息表 + 步骤说明 + 步骤执行表 + 测量数据表 + 审计记录 + 结论）
- Excel 数据表（测量原始数据，多指标分 sheet）

文件存 backend/storage/reports/{experiment_id}/，幂等（已生成则复用）。
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from loguru import logger
from sqlalchemy import select

STORAGE_ROOT = Path(__file__).resolve().parents[3] / "storage" / "reports"


def _build_report_dir(experiment_id: str) -> Path:
    d = STORAGE_ROOT / experiment_id
    d.mkdir(parents=True, exist_ok=True)
    return d


async def generate_report(experiment_id: str) -> dict[str, str]:
    """生成实验报告（Word + Excel），返回文件路径。

    幂等：若 Word 已存在则直接复用。
    """
    from app.database.session import get_session_factory
    from app.models.tables import (
        Experiment,
        ExperimentStep,
        Measurement,
        ExperimentAudit,
        Protocol,
    )

    report_dir = _build_report_dir(experiment_id)
    word_path = report_dir / "报告.docx"
    excel_path = report_dir / "数据表.xlsx"

    if word_path.exists() and excel_path.exists():
        logger.bind(component="report").info("报告已存在，复用: {}", experiment_id)
        return {"word_path": str(word_path), "excel_path": str(excel_path)}

    factory = get_session_factory()
    async with factory() as session:
        exp = await session.get(Experiment, experiment_id)
        if exp is None:
            raise ValueError(f"实验不存在: {experiment_id}")

        protocol = await session.get(Protocol, exp.protocol_id) if exp.protocol_id else None

        steps = (
            await session.execute(
                select(ExperimentStep)
                .where(ExperimentStep.experiment_id == experiment_id)
                .order_by(ExperimentStep.step_order.asc())
            )
        ).scalars().all()

        measurements = (
            await session.execute(
                select(Measurement)
                .where(Measurement.experiment_id == experiment_id)
                .order_by(Measurement.timestamp.asc())
            )
        ).scalars().all()

        audits = (
            await session.execute(
                select(ExperimentAudit)
                .where(ExperimentAudit.experiment_id == experiment_id)
                .order_by(ExperimentAudit.created_at.asc())
            )
        ).scalars().all()

    # 解析 result 里的摘要（若有）
    summary: dict[str, Any] = {}
    if exp.result:
        try:
            result_data = json.loads(exp.result)
            summary = result_data.get("summary", {})
        except (json.JSONDecodeError, AttributeError):
            pass

    # 生成 Word
    _gen_word(
        word_path,
        experiment_id,
        exp,
        protocol,
        steps,
        measurements,
        audits,
        summary,
    )

    # 生成 Excel
    _gen_excel(excel_path, experiment_id, measurements)

    logger.bind(component="report").info(
        "报告生成完成: {} -> {}", experiment_id, word_path
    )
    return {"word_path": str(word_path), "excel_path": str(excel_path)}


def _gen_word(
    path: Path,
    experiment_id: str,
    exp: Any,
    protocol: Any,
    steps: list[Any],
    measurements: list[Any],
    audits: list[Any],
    summary: dict[str, Any],
) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("实验报告", level=0)
    doc.add_heading(exp.name or experiment_id, level=1)

    # 实验信息表
    doc.add_heading("一、实验信息", level=2)
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    info_rows = [
        ("实验编号", experiment_id),
        ("实验名称", exp.name or ""),
        ("操作员", exp.operator or ""),
        ("方案", protocol.name if protocol else (exp.protocol_id or "")),
        ("样品编号", exp.sample_code or ""),
        ("状态", exp.status or ""),
        ("创建时间", str(exp.created_at or "")),
    ]
    for key, val in info_rows:
        row = info_table.add_row()
        row.cells[0].text = key
        row.cells[1].text = str(val)

    # 方案步骤说明
    if protocol:
        doc.add_heading("二、方案步骤", level=2)
        for s in steps:
            params = s.params or "{}"
            doc.add_paragraph(
                f"步骤 {s.step_order}：{s.action}（{params}）— {s.status}",
                style="List Number",
            )

    # 步骤执行表
    doc.add_heading("三、步骤执行明细", level=2)
    step_table = doc.add_table(rows=1, cols=5)
    step_table.style = "Table Grid"
    hdr = step_table.rows[0].cells
    for i, h in enumerate(["序号", "设备", "动作", "状态", "错误"]):
        hdr[i].text = h
    for s in steps:
        row = step_table.add_row().cells
        row[0].text = str(s.step_order)
        row[1].text = s.device_id
        row[2].text = s.action
        row[3].text = s.status
        row[4].text = s.error_message or ""

    # 测量数据表
    doc.add_heading("四、测量数据", level=2)
    if summary:
        doc.add_paragraph(
            f"指标 {summary.get('metric_name','')}：共 {summary.get('points',0)} 点，"
            f"峰值 {summary.get('max','')}，终值 {summary.get('final','')}"
        )
    if measurements:
        meas_table = doc.add_table(rows=1, cols=4)
        meas_table.style = "Table Grid"
        hdr = meas_table.rows[0].cells
        for i, h in enumerate(["时间", "指标", "值", "单位"]):
            hdr[i].text = h
        for m in measurements[:100]:  # 上限 100 行，避免报告过长
            row = meas_table.add_row().cells
            row[0].text = str(m.timestamp or "")
            row[1].text = m.metric_name
            row[2].text = str(m.metric_value)
            row[3].text = m.unit or ""

    # 审计记录
    doc.add_heading("五、执行记录", level=2)
    for a in audits:
        doc.add_paragraph(f"[{a.created_at}] {a.event_type}: {a.detail}")

    # 结论
    doc.add_heading("六、结论", level=2)
    if summary:
        doc.add_paragraph(
            f"本实验共采集 {summary.get('points',0)} 个数据点，"
            f"指标 {summary.get('metric_name','')} 范围为 "
            f"{summary.get('min','')} ~ {summary.get('max','')}，"
            f"最终值为 {summary.get('final','')}。"
        )
    else:
        doc.add_paragraph("实验执行完成，数据见测量数据表。")

    doc.save(str(path))


def _gen_excel(path: Path, experiment_id: str, measurements: list[Any]) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    # 按指标分组，每个指标一个 sheet
    by_metric: dict[str, list[Any]] = {}
    for m in measurements:
        by_metric.setdefault(m.metric_name, []).append(m)

    if not by_metric:
        by_metric = {"(空)": []}

    for idx, (metric, rows) in enumerate(by_metric.items()):
        ws = wb.active if idx == 0 else wb.create_sheet()
        ws.title = metric[:31]  # sheet 名上限 31 字符
        ws.append(["时间戳", "指标", "值", "单位"])
        for m in rows:
            ws.append([str(m.timestamp or ""), m.metric_name, m.metric_value, m.unit or ""])

    wb.save(str(path))


__all__ = ["generate_report"]
